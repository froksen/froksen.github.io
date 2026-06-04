#!/usr/bin/env python3
"""
Genererer en Word-udgave af naturfagsprøve-siden.
- Klikbare hyperlinks på alle ressourcer (SIM, VIDEO, DR)
- SDG-billeder indlejret som inline-ikoner med tekst
"""

import os, re
from bs4 import BeautifulSoup
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT
from docx.enum import text as docx_breaks

SDG_CACHE = '/home/olefrandsen/Udvikling/froksen.github.io/fff/sdg_cache'

# ─────────────────────────────────────────────
# Hjælpefunktioner – formatering
# ─────────────────────────────────────────────

def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)

def remove_table_borders(tbl):
    for row in tbl.rows:
        for c in row.cells:
            tc = c._tc
            tcPr = tc.get_or_add_tcPr()
            tcBorders = OxmlElement('w:tcBorders')
            for side in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
                el = OxmlElement(f'w:{side}')
                el.set(qn('w:val'), 'none')
                tcBorders.append(el)
            tcPr.append(tcBorders)

def add_horizontal_rule(doc):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(5)
    p.paragraph_format.space_after = Pt(5)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'AAAAAA')
    pBdr.append(bottom)
    pPr.append(pBdr)

def page_break(doc):
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(docx_breaks.WD_BREAK.PAGE)

# ─────────────────────────────────────────────
# Hyperlink-hjælper
# ─────────────────────────────────────────────

def add_hyperlink(paragraph, text, url, font_size=9, bold=False, color_hex=None):
    """Indsæt klikbart hyperlink i et afsnit."""
    part = paragraph.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hl = OxmlElement('w:hyperlink')
    hl.set(qn('r:id'), r_id)

    run = OxmlElement('w:r')
    rPr = OxmlElement('w:rPr')

    # Understreget blå hyperlink-stil
    rStyle = OxmlElement('w:rStyle')
    rStyle.set(qn('w:val'), 'Hyperlink')
    rPr.append(rStyle)

    sz = OxmlElement('w:sz')
    sz.set(qn('w:val'), str(font_size * 2))
    rPr.append(sz)
    szCs = OxmlElement('w:szCs')
    szCs.set(qn('w:val'), str(font_size * 2))
    rPr.append(szCs)

    if bold:
        rPr.append(OxmlElement('w:b'))

    if color_hex:
        clr = OxmlElement('w:color')
        clr.set(qn('w:val'), color_hex)
        rPr.append(clr)

    run.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    t.set('{http://www.w3.org/XML/1998/namespace}space', 'preserve')
    run.append(t)
    hl.append(run)
    paragraph._p.append(hl)

# ─────────────────────────────────────────────
# SDG-billedhjælper
# ─────────────────────────────────────────────

def goal_img_path(src_url):
    """Returner lokal sti til SDG-billede ud fra URL."""
    m = re.search(r'goal-(\d+)\.png', src_url)
    if m:
        path = os.path.join(SDG_CACHE, f'goal-{m.group(1)}.png')
        if os.path.exists(path):
            return path
    return None

def add_sdg_images(doc, sdg_goals):
    """
    sdg_goals = liste af (img_src, caption_text)
    Indsætter en lille tabel med SDG-ikoner (Cm(1.4)) og tekst under.
    """
    if not sdg_goals:
        return

    n = len(sdg_goals)
    tbl = doc.add_table(rows=2, cols=n)
    tbl.style = 'Table Grid'
    remove_table_borders(tbl)

    # Overskrift
    label_p = doc.paragraphs[-1]  # intet – vi tilføjer label manuelt

    p_lbl = doc.add_paragraph()
    # Vi indsætter tabellen, ikke p_lbl – fjern den igen bagefter
    # Faktisk: vi tilføjer labeltekst i første omgang og tabel derunder
    # Lad os gøre det ordentligt: tilføj label FØR tabellen ved at lave
    # tabellen direkte som næste element.

    # Fjern den tomme paragraf vi netop tilføjede
    lbl_el = p_lbl._element
    lbl_el.getparent().remove(lbl_el)

    # Tilføj label-paragraf inden tabellen
    p_lbl2 = doc.add_paragraph()
    p_lbl2.paragraph_format.space_before = Pt(6)
    p_lbl2.paragraph_format.space_after = Pt(2)
    r_lbl = p_lbl2.add_run("FN's Verdensmål")
    r_lbl.font.bold = True
    r_lbl.font.size = Pt(9)
    r_lbl.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)

    tbl = doc.add_table(rows=2, cols=n)
    tbl.style = 'Table Grid'
    remove_table_borders(tbl)

    for i, (src, caption) in enumerate(sdg_goals):
        # Billed-celle (øverste række)
        img_cell = tbl.rows[0].cells[i]
        img_cell.width = Cm(1.8)
        p_img = img_cell.paragraphs[0]
        p_img.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_img.paragraph_format.space_before = Pt(0)
        p_img.paragraph_format.space_after = Pt(0)

        local = goal_img_path(src)
        if local:
            run_img = p_img.add_run()
            run_img.add_picture(local, width=Cm(1.4))
        else:
            run_img = p_img.add_run('[?]')

        # Tekst-celle (nederste række)
        txt_cell = tbl.rows[1].cells[i]
        p_txt = txt_cell.paragraphs[0]
        p_txt.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p_txt.paragraph_format.space_before = Pt(1)
        p_txt.paragraph_format.space_after = Pt(2)
        r_txt = p_txt.add_run(caption)
        r_txt.font.size = Pt(7)
        r_txt.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

    # Lille mellemrum efter tabellen
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(2)

# ─────────────────────────────────────────────
# Parse HTML
# ─────────────────────────────────────────────

with open('/home/olefrandsen/Udvikling/froksen.github.io/fff/index.html', 'r', encoding='utf-8') as f:
    soup = BeautifulSoup(f.read(), 'html.parser')

# ─────────────────────────────────────────────
# Opret dokument
# ─────────────────────────────────────────────

doc = Document()

for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

normal = doc.styles['Normal']
normal.font.name = 'Calibri'
normal.font.size = Pt(10)

# ─────────────────────────────────────────────
# Genbrugelige bygge-funktioner
# ─────────────────────────────────────────────

def add_section_header(doc, text):
    tbl = doc.add_table(rows=1, cols=1)
    tbl.style = 'Table Grid'
    cell = tbl.rows[0].cells[0]
    set_cell_bg(cell, '4A5FC1')
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    remove_table_borders(tbl)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)

def add_article_header(doc, text):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '8')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), '667EEA')
    pBdr.append(bottom)
    pPr.append(pBdr)

def add_subsection_label(doc, text, color_hex='2C5F7F'):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(text)
    run.font.bold = True
    run.font.size = Pt(10)
    r, g, b = int(color_hex[0:2],16), int(color_hex[2:4],16), int(color_hex[4:6],16)
    run.font.color.rgb = RGBColor(r, g, b)

def add_bullet(doc, text, bold_part=None):
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)
    if bold_part:
        rb = p.add_run(bold_part + ': ')
        rb.font.bold = True
        rb.font.size = Pt(10)
    r = p.add_run(text)
    r.font.size = Pt(10)

def add_focus_table(doc, fysik_items, geo_items, bio_items):
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = 'Table Grid'
    headers = [
        ('Fysik-kemi', 'FFD0D0', 'C92A2A', fysik_items),
        ('Geografi',   'D0F0D8', '2B8A3E', geo_items),
        ('Biologi',    'E8D8FF', '5C3DC9', bio_items),
    ]
    for i, (label, bg, fg, items) in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        set_cell_bg(cell, bg)
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(label)
        run.font.bold = True
        run.font.size = Pt(10)
        r,g,b = int(fg[0:2],16), int(fg[2:4],16), int(fg[4:6],16)
        run.font.color.rgb = RGBColor(r,g,b)
        for item in items:
            bp = cell.add_paragraph()
            bp.paragraph_format.space_before = Pt(0)
            bp.paragraph_format.space_after = Pt(1)
            bp.paragraph_format.left_indent = Cm(0.3)
            bp.add_run('• ' + item).font.size = Pt(9)
    for cell in tbl.rows[0].cells:
        cell.width = Cm(5.3)
    doc.add_paragraph().paragraph_format.space_after = Pt(3)

def _parse_forsog(details_el):
    """Udtræk forsøgsbeskrivelse fra <details class='forsog-detaljer'>."""
    div = details_el.find('div', class_='forsog-beskrivelse')
    if not div:
        return None
    data = {}
    for p_el in div.find_all('p', recursive=False):
        strong = p_el.find('strong')
        if not strong:
            continue
        key = strong.get_text(strip=True).rstrip(':')
        strong.extract()
        value = p_el.get_text(separator=' ', strip=True).lstrip(':').strip()
        data[key] = value
    ol = div.find('ol')
    if ol:
        data['_steps'] = [li.get_text(strip=True) for li in ol.find_all('li')]
    return data


def add_forsog_beskrivelse(doc, forsog):
    """Tilføj forsøgsbeskrivelse som indrykkede underafsnit."""
    if not forsog:
        return
    for key in ['Formål', 'Materialer', 'Fremgangsmåde', 'Databehandling', 'Diskussion']:
        if key == 'Fremgangsmåde' and '_steps' in forsog:
            p_hdr = doc.add_paragraph()
            p_hdr.paragraph_format.space_before = Pt(1)
            p_hdr.paragraph_format.space_after = Pt(0)
            p_hdr.paragraph_format.left_indent = Cm(1.0)
            rb = p_hdr.add_run('Fremgangsmåde:')
            rb.font.bold = True
            rb.font.size = Pt(8)
            rb.font.color.rgb = RGBColor(0x2C, 0x5F, 0x7F)
            for i, step in enumerate(forsog['_steps'], 1):
                ps = doc.add_paragraph()
                ps.paragraph_format.space_before = Pt(0)
                ps.paragraph_format.space_after = Pt(0)
                ps.paragraph_format.left_indent = Cm(1.5)
                ps.add_run(f'{i}. {step}').font.size = Pt(8)
        elif key in forsog and forsog[key]:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.left_indent = Cm(1.0)
            rb = p.add_run(key + ': ')
            rb.font.bold = True
            rb.font.size = Pt(8)
            rb.font.color.rgb = RGBColor(0x2C, 0x5F, 0x7F)
            p.add_run(forsog[key]).font.size = Pt(8)
    # Lille luft efter beskrivelsen
    sp = doc.add_paragraph()
    sp.paragraph_format.space_after = Pt(3)


def add_resource_line(doc, tag, link_text, url, description):
    """Ressource-linje med klikbart hyperlink."""
    p = doc.add_paragraph(style='List Bullet')
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.left_indent = Cm(0.5)

    if tag in ('SIMULERING', 'SIM'):
        tag_color = RGBColor(0x15, 0x65, 0xC0)
        tag_label = '[SIM] '
    elif tag == 'DR':
        tag_color = RGBColor(0x8B, 0x00, 0x00)
        tag_label = '[DR] '
    else:
        tag_color = RGBColor(0xC6, 0x28, 0x28)
        tag_label = '[VIDEO] '

    r1 = p.add_run(tag_label)
    r1.font.bold = True
    r1.font.size = Pt(9)
    r1.font.color.rgb = tag_color

    if url:
        add_hyperlink(p, link_text, url, font_size=9, bold=True)
    else:
        r2 = p.add_run(link_text)
        r2.font.bold = True
        r2.font.size = Pt(9)

    if description:
        r3 = p.add_run(' – ' + description)
        r3.font.size = Pt(9)

# ─────────────────────────────────────────────
# Parse artikelindhold (inkl. SDG-billeder og link-URLs)
# ─────────────────────────────────────────────

def parse_article(article):
    result = {}

    h3 = article.find('h3')
    result['title'] = h3.get_text(strip=True) if h3 else ''

    # Undersøgelser
    result['investigations'] = []
    inv_div = article.find('div', class_='investigations')
    if inv_div:
        for li in inv_div.find_all('li'):
            strong = li.find('strong')
            bold = strong.get_text(strip=True).rstrip(':') if strong else None
            if strong:
                strong.extract()
            # Udtræk og fjern <details> inden get_text()
            details_el = li.find('details', class_='forsog-detaljer')
            forsog = None
            if details_el:
                forsog = _parse_forsog(details_el)
                details_el.extract()
            rest = li.get_text(separator=' ', strip=True).strip().lstrip(':').strip()
            result['investigations'].append((bold, rest, forsog))

    # Ressourcer – nu med URL
    result['resources'] = []
    res_div = article.find('div', class_='resources')
    if res_div:
        for li in res_div.find_all('li'):
            tag_span = li.find('span', class_='resource-tag')
            tag = tag_span.get_text(strip=True) if tag_span else 'VIDEO'
            if tag_span:
                tag_span.extract()
            link = li.find('a')
            link_text = link.get_text(strip=True) if link else ''
            url = link.get('href', '') if link else ''
            if link:
                link.extract()
            desc = li.get_text(separator=' ', strip=True).strip(' –-').strip()
            result['resources'].append((tag, link_text, url, desc))

    # Fokusområder
    result['fysik'] = []
    result['geografi'] = []
    result['biologi'] = []
    focus_div = article.find('div', class_='focus-areas')
    if focus_div:
        for col in focus_div.find_all('div', class_='focus-column'):
            classes = col.get('class', [])
            items = [li.get_text(strip=True) for li in col.find_all('li')]
            if 'fysik' in classes:
                result['fysik'] = items
            elif 'geografi' in classes:
                result['geografi'] = items
            elif 'biologi' in classes:
                result['biologi'] = items

    # SDG – nu med billede-URL
    result['sdg'] = []
    sdg_div = article.find('div', class_='sdg-section')
    if sdg_div:
        for goal_div in sdg_div.find_all('div', class_='sdg-goal'):
            img = goal_div.find('img')
            p_el = goal_div.find('p')
            if img and p_el:
                result['sdg'].append((img.get('src', ''), p_el.get_text(strip=True)))

    return result

# ─────────────────────────────────────────────
# TOC – første side
# ─────────────────────────────────────────────

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
p.paragraph_format.space_before = Pt(20)
p.paragraph_format.space_after = Pt(4)
run = p.add_run('Den fællesfaglige naturfagsprøve i 9. klasse')
run.font.bold = True
run.font.size = Pt(22)
run.font.color.rgb = RGBColor(0x4A, 0x5F, 0xC1)

p2 = doc.add_paragraph()
p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
p2.paragraph_format.space_after = Pt(20)
run2 = p2.add_run('Idéer til delemner, undersøgelser og modeller')
run2.font.size = Pt(12)
run2.font.italic = True
run2.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

add_horizontal_rule(doc)

p_toc = doc.add_paragraph()
p_toc.paragraph_format.space_before = Pt(10)
p_toc.paragraph_format.space_after = Pt(8)
r_toc = p_toc.add_run('Indholdsfortegnelse')
r_toc.font.bold = True
r_toc.font.size = Pt(14)
r_toc.font.color.rgb = RGBColor(0x4A, 0x5F, 0xC1)

theme_sections = soup.find_all('section', class_='theme-section')
model_section  = soup.find('section', id='modeller')

toc_data = []
for section in theme_sections:
    hdr = section.find('div', class_='theme-header')
    section_title = hdr.get_text(strip=True) if hdr else '?'
    articles = [art.find('h3').get_text(strip=True)
                for art in section.find_all('article', class_='subtopic')
                if art.find('h3')]
    toc_data.append((section_title, articles))

model_cards_list = [c.find('h3').get_text(strip=True)
                    for c in (model_section.find_all('div', class_='model-card') if model_section else [])
                    if c.find('h3')]

# TOC-tabel: 2 kolonner
toc_tbl = doc.add_table(rows=1, cols=2)
toc_tbl.style = 'Table Grid'
remove_table_borders(toc_tbl)

def fill_toc_column(cell, sections_list):
    first_sec = True
    for section_title, articles in sections_list:
        if not first_sec:
            cell.add_paragraph().paragraph_format.space_after = Pt(4)
        first_sec = False
        hp = cell.add_paragraph()
        hp.paragraph_format.space_before = Pt(0)
        hp.paragraph_format.space_after = Pt(2)
        hr = hp.add_run(section_title)
        hr.font.bold = True
        hr.font.size = Pt(11)
        hr.font.color.rgb = RGBColor(0x4A, 0x5F, 0xC1)
        for i, art_title in enumerate(articles, 1):
            ap = cell.add_paragraph()
            ap.paragraph_format.space_before = Pt(0)
            ap.paragraph_format.space_after = Pt(1)
            ap.paragraph_format.left_indent = Cm(0.5)
            ap.add_run(f'{i}. {art_title}').font.size = Pt(9)

left_cell  = toc_tbl.rows[0].cells[0]
right_cell = toc_tbl.rows[0].cells[1]
fill_toc_column(left_cell, toc_data[:2])
fill_toc_column(right_cell, toc_data[2:])

# Modeller i højre kolonne
right_cell.add_paragraph().paragraph_format.space_after = Pt(4)
hp2 = right_cell.add_paragraph()
hp2.paragraph_format.space_before = Pt(0)
hp2.paragraph_format.space_after = Pt(2)
hp2.add_run('Modeller').font.bold = True
hp2.runs[0].font.size = Pt(11)
hp2.runs[0].font.color.rgb = RGBColor(0x4A, 0x5F, 0xC1)
for i, mc in enumerate(model_cards_list, 1):
    ap = right_cell.add_paragraph()
    ap.paragraph_format.space_before = Pt(0)
    ap.paragraph_format.space_after = Pt(1)
    ap.paragraph_format.left_indent = Cm(0.5)
    ap.add_run(f'{i}. {mc}').font.size = Pt(9)

add_horizontal_rule(doc)
p_note = doc.add_paragraph()
p_note.paragraph_format.space_before = Pt(4)
r_note = p_note.add_run(
    'Dokumentet er organiseret med sideskift mellem hvert fordybelsesområde. '
    'Hvert delemne indeholder forslag til elevundersøgelser, klikbare digitale ressourcer, '
    "faglige fokusområder (fysik-kemi, geografi, biologi) og FN's Verdensmål med ikoner."
)
r_note.font.size = Pt(9)
r_note.font.italic = True
r_note.font.color.rgb = RGBColor(0x77, 0x77, 0x77)

page_break(doc)

# ─────────────────────────────────────────────
# Fordybelsesområder
# ─────────────────────────────────────────────

for sec_idx, section in enumerate(theme_sections):
    if sec_idx > 0:
        page_break(doc)

    hdr = section.find('div', class_='theme-header')
    add_section_header(doc, hdr.get_text(strip=True) if hdr else '?')

    articles = section.find_all('article', class_='subtopic')
    for art_idx, article in enumerate(articles):
        data = parse_article(article)

        add_article_header(doc, data['title'])

        if data['investigations']:
            add_subsection_label(doc, 'Undersøgelser elever kan lave', '1A6A9A')
            for bold, text, forsog in data['investigations']:
                add_bullet(doc, text, bold_part=bold)
                add_forsog_beskrivelse(doc, forsog)

        if data['resources']:
            add_subsection_label(doc, 'Digitale ressourcer', 'B8860B')
            for tag, link_text, url, desc in data['resources']:
                add_resource_line(doc, tag, link_text, url, desc)

        if data['fysik'] or data['geografi'] or data['biologi']:
            add_subsection_label(doc, 'Faglige fokusområder', '444444')
            add_focus_table(doc, data['fysik'], data['geografi'], data['biologi'])

        if data['sdg']:
            add_sdg_images(doc, data['sdg'])

        if art_idx < len(articles) - 1:
            add_horizontal_rule(doc)

# ─────────────────────────────────────────────
# Modeller
# ─────────────────────────────────────────────

page_break(doc)
add_section_header(doc, 'Modeller')

models_intro = model_section.find('p', class_='models-intro') if model_section else None
if models_intro:
    p_intro = doc.add_paragraph()
    p_intro.paragraph_format.space_after = Pt(10)
    r_intro = p_intro.add_run(models_intro.get_text(strip=True))
    r_intro.font.size = Pt(10)
    r_intro.font.italic = True

model_cards_divs = model_section.find_all('div', class_='model-card') if model_section else []
for i, card in enumerate(model_cards_divs):
    h3       = card.find('h3')
    relevance = card.find('p', class_='model-relevance')
    desc_el  = card.find('p', class_='model-desc')
    link_el  = card.find('a', class_='model-link')

    if h3:
        p_h = doc.add_paragraph()
        p_h.paragraph_format.space_before = Pt(8)
        p_h.paragraph_format.space_after = Pt(2)
        r_h = p_h.add_run(h3.get_text(strip=True))
        r_h.font.bold = True
        r_h.font.size = Pt(11)
        r_h.font.color.rgb = RGBColor(0x76, 0x4B, 0xA2)

    if relevance:
        p_r = doc.add_paragraph()
        p_r.paragraph_format.space_before = Pt(0)
        p_r.paragraph_format.space_after = Pt(2)
        r_r = p_r.add_run(relevance.get_text(strip=True))
        r_r.font.size = Pt(9)
        r_r.font.italic = True
        r_r.font.color.rgb = RGBColor(0x66, 0x7E, 0xEA)

    if desc_el:
        p_d = doc.add_paragraph()
        p_d.paragraph_format.space_before = Pt(0)
        p_d.paragraph_format.space_after = Pt(2)
        p_d.add_run(desc_el.get_text(strip=True)).font.size = Pt(10)

    if link_el:
        url = link_el.get('href', '')
        link_txt = link_el.get_text(strip=True)
        p_l = doc.add_paragraph()
        p_l.paragraph_format.space_before = Pt(0)
        p_l.paragraph_format.space_after = Pt(2)
        r_ll = p_l.add_run('Link: ')
        r_ll.font.size = Pt(9)
        r_ll.font.bold = True
        if url:
            add_hyperlink(p_l, link_txt, url, font_size=9)
        else:
            p_l.add_run(link_txt).font.size = Pt(9)

    if i < len(model_cards_divs) - 1:
        add_horizontal_rule(doc)

# ─────────────────────────────────────────────
# Gem
# ─────────────────────────────────────────────

out_path = '/home/olefrandsen/Udvikling/froksen.github.io/fff/naturfagsprove.docx'
doc.save(out_path)
print(f'Gemt: {out_path}')
